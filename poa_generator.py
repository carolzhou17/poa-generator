"""
poa_generator.py  —  Pure logic, no UI.
Handles field definitions, Word replacement, CSV parsing, and document generation.
Supports: 1-IP/1-agent (Shi & Aispuro style) and 2-IP/3-agent (Ding & Luo style).
"""

import io
import csv
from typing import Optional, List
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Field definitions
# ---------------------------------------------------------------------------
# per_principal=True  → belongs to one principal (repeated for each IP)
# per_principal=False → shared across the whole case (surrogate, agent, etc.)

FIELDS = [
    # ---- Per-principal fields ----
    {
        "key":           "principal_name",
        "label":         "Full Name",
        "section":       "Principal",
        "template_value":"CHENGFANG SHI",
        "required":      True,
        "to_upper":      True,
        "default":       "",
        "help":          "ALL CAPS — e.g., JOHN SMITH",
        "per_principal": True,
    },
    {
        "key":           "principal_role",
        "label":         "Role",
        "section":       "Principal",
        "template_value":"Intended Father",
        "required":      False,
        "to_upper":      False,
        "default":       "Intended Father",
        "help":          "Intended Father  or  Intended Mother",
        "per_principal": True,
    },
    {
        "key":           "passport_country",
        "label":         "Passport Country",
        "section":       "Principal",
        "template_value":"People’s Republic of China",
        "required":      False,
        "to_upper":      False,
        "default":       "People’s Republic of China",
        "help":          "Country as it appears on passport",
        "per_principal": True,
    },
    {
        "key":           "principal_passport",
        "label":         "Passport Number",
        "section":       "Principal",
        "template_value":"ED3297013",
        "required":      True,
        "to_upper":      False,
        "default":       "",
        "help":          "e.g., ED3297013",
        "per_principal": True,
    },

    # ---- Shared case fields ----
    {
        "key":           "child_last_name",
        "label":         "Child Last Name",
        "section":       "Child",
        "template_value":"SHI",
        "required":      True,
        "to_upper":      True,
        "default":       "",
        "help":          "Will appear as 'infant CHILD [LAST NAME]'",
        "per_principal": False,
    },
    {
        "key":           "surrogate_name",
        "label":         "Surrogate Full Name",
        "section":       "Surrogate",
        "template_value":"SELENA MARIA AISPURO",
        "required":      True,
        "to_upper":      True,
        "default":       "",
        "help":          "ALL CAPS",
        "per_principal": False,
    },
    {
        "key":           "surrogate_dob",
        "label":         "Surrogate Date of Birth",
        "section":       "Surrogate",
        "template_value":"May 22, 1996",
        "required":      True,
        "to_upper":      False,
        "default":       "",
        "help":          "e.g., May 22, 1996",
        "per_principal": False,
    },
    {
        "key":           "due_date",
        "label":         "Due Date",
        "section":       "Birth Details",
        "template_value":"December 4, 2025",
        "required":      True,
        "to_upper":      False,
        "default":       "",
        "help":          "e.g., March 15, 2026",
        "per_principal": False,
    },
    {
        "key":           "hospital_name",
        "label":         "Hospital Name",
        "section":       "Birth Details",
        "template_value":"Loma Linda University Medical Center",
        "required":      False,
        "to_upper":      False,
        "default":       "Loma Linda University Medical Center",
        "help":          "Full hospital name",
        "per_principal": False,
    },
    {
        "key":           "hospital_city",
        "label":         "Hospital City",
        "section":       "Birth Details",
        "template_value":"Loma Linda",
        "required":      False,
        "to_upper":      False,
        "default":       "Loma Linda",
        "help":          "",
        "per_principal": False,
    },
    {
        "key":           "hospital_state",
        "label":         "Hospital State",
        "section":       "Birth Details",
        "template_value":"California",
        "required":      False,
        "to_upper":      False,
        "default":       "California",
        "help":          "",
        "per_principal": False,
    },
    {
        "key":           "agent_name",
        "label":         "Agent / Attorney-in-Fact Full Name",
        "section":       "Agent",
        "template_value":"JIAJIA GAO",
        "required":      True,
        "to_upper":      True,
        "default":       "",
        "help":          "ALL CAPS",
        "per_principal": False,
    },
    {
        "key":           "agent_dob",
        "label":         "Agent Date of Birth",
        "section":       "Agent",
        "template_value":"02/24/1986",
        "required":      True,
        "to_upper":      False,
        "default":       "",
        "help":          "e.g., 02/24/1986",
        "per_principal": False,
    },
    {
        "key":           "agent_passport",
        "label":         "Agent Passport Number",
        "section":       "Agent",
        "template_value":"EJ4979964",
        "required":      True,
        "to_upper":      False,
        "default":       "",
        "help":          "e.g., EJ4979964",
        "per_principal": False,
    },
    {
        "key":           "agent_pronoun",
        "label":         "Agent Pronoun",
        "section":       "Agent",
        "template_value":"her",
        "required":      False,
        "to_upper":      False,
        "default":       "her",
        "help":          "her / his / their",
        "per_principal": False,
    },
    {
        "key":           "attorney_name",
        "label":         "Handling Attorney",
        "section":       "Firm",
        "template_value":"Xuelan Fang",
        "required":      False,
        "to_upper":      False,
        "default":       "Xuelan Fang",
        "help":          "",
        "per_principal": False,
    },
    {
        "key":           "agency_name",
        "label":         "Fertility Agency (for filename)",
        "section":       "Firm",
        "template_value":"",
        "required":      False,
        "to_upper":      False,
        "default":       "",
        "help":          "e.g., C&T Fertility Consultant",
        "per_principal": False,
    },
]

PRINCIPAL_FIELDS = [f for f in FIELDS if f["per_principal"]]
CASE_FIELDS      = [f for f in FIELDS if not f["per_principal"]]

# Replacement order: specific strings before the substrings they contain.
REPLACEMENT_ORDER = [
    "hospital_name",       # "Loma Linda University Medical Center" before "Loma Linda"
    "hospital_city",
    "hospital_state",
    "due_date",
    "surrogate_dob",
    "agent_dob",
    "principal_name",
    "child_last_name",     # special: replaces "CHILD SHI", not just "SHI"
    "surrogate_name",
    "agent_name",
    "attorney_name",
    "principal_passport",
    "agent_passport",
    "passport_country",
    "principal_role",
    "agent_pronoun",
]


# ---------------------------------------------------------------------------
# Low-level Word replacement engine
# ---------------------------------------------------------------------------

def _replace_once(para, old: str, new: str) -> bool:
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    idx, end_idx = full.index(old), full.index(old) + len(old)
    pos = 0
    s_run = s_off = e_run = e_off = None
    for i, run in enumerate(para.runs):
        run_end = pos + len(run.text)
        if s_run is None and run_end > idx:
            s_run, s_off = i, idx - pos
        if e_run is None and run_end >= end_idx:
            e_run, e_off = i, end_idx - pos
            break
        pos += len(run.text)
    if s_run is None or e_run is None:
        return False
    runs = list(para.runs)
    if s_run == e_run:
        r = runs[s_run]
        r.text = r.text[:s_off] + new + r.text[e_off:]
    else:
        runs[s_run].text = runs[s_run].text[:s_off] + new
        for i in range(s_run + 1, e_run):
            runs[i].text = ""
        runs[e_run].text = runs[e_run].text[e_off:]
    return True


def _replace_in_paragraph(para, old: str, new: str) -> int:
    count = 0
    while _replace_once(para, old, new):
        count += 1
    return count


def _replace_in_doc(doc, old: str, new: str) -> int:
    if not old or old == new:
        return 0
    count = 0
    for para in doc.paragraphs:
        count += _replace_in_paragraph(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += _replace_in_paragraph(para, old, new)
    return count


# ---------------------------------------------------------------------------
# Image replacement
# ---------------------------------------------------------------------------

def _find_image_paragraphs(doc) -> list:
    return [p for p in doc.paragraphs
            if p._element.find('.//' + qn('a:blip')) is not None]


def _replace_image(doc, para, new_bytes: bytes) -> bool:
    blip = para._element.find('.//' + qn('a:blip'))
    if blip is None:
        return False
    rid = blip.get(qn('r:embed'))
    if rid is None:
        return False
    doc.part.related_parts[rid]._blob = new_bytes
    return True


# ---------------------------------------------------------------------------
# Document merging
# ---------------------------------------------------------------------------

def _merge_documents(doc_bytes_list: List[bytes]) -> bytes:
    """Combine multiple .docx files into one, each starting on a new page."""
    if len(doc_bytes_list) == 1:
        return doc_bytes_list[0]

    from docxcompose.composer import Composer

    base = Document(io.BytesIO(doc_bytes_list[0]))
    composer = Composer(base)

    for extra_bytes in doc_bytes_list[1:]:
        src = Document(io.BytesIO(extra_bytes))
        composer.append(src)

    out = io.BytesIO()
    composer.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def validate_case(case_info: dict, principals: Optional[List[dict]] = None) -> list:
    """Return list of error strings. Empty = valid."""
    errors = []
    # Validate shared case fields
    for field in CASE_FIELDS:
        if field["required"] and not case_info.get(field["key"], "").strip():
            errors.append(f"'{field['label']}' is required.")
    # Validate each principal
    if principals:
        for i, p in enumerate(principals, 1):
            for field in PRINCIPAL_FIELDS:
                if field["required"] and not p.get(field["key"], "").strip():
                    errors.append(f"Principal {i}: '{field['label']}' is required.")
    return errors


def make_filename(case_info: dict, principals: Optional[List[dict]] = None) -> str:
    if principals:
        last_names = [p.get("principal_name", "Client").split()[-1] for p in principals]
    else:
        last_names = [case_info.get("principal_name", "Client").split()[-1]]

    surrogate_last = case_info.get("surrogate_name", "Surrogate").split()[-1]
    agency = case_info.get("agency_name", "").strip()

    principals_str = " & ".join(last_names)
    parts = [f"POA - {principals_str} & {surrogate_last}"]
    if agency:
        parts.append(agency)
    return " - ".join(parts) + ".docx"


def generate_poa_bytes(
    template_bytes: bytes,
    info: dict,
    principal_photo: Optional[bytes] = None,
    agent_photo: Optional[bytes] = None,
) -> tuple:
    """Generate one POA document (single principal). Returns (bytes, count)."""
    doc = Document(io.BytesIO(template_bytes))
    total = 0

    for key in REPLACEMENT_ORDER:
        field = next((f for f in FIELDS if f["key"] == key), None)
        if field is None:
            continue
        value = info.get(key, "").strip()
        if field["to_upper"]:
            value = value.upper()
        if not value:
            value = field["default"]
        template_val = field["template_value"]

        if key == "child_last_name":
            total += _replace_in_doc(doc, f"CHILD {template_val}", f"CHILD {value}")
        else:
            total += _replace_in_doc(doc, template_val, value)

    img_paras = _find_image_paragraphs(doc)
    if principal_photo and len(img_paras) > 0:
        _replace_image(doc, img_paras[0], principal_photo)
    if agent_photo and len(img_paras) > 1:
        _replace_image(doc, img_paras[1], agent_photo)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), total


def generate_multi_poa_bytes(
    template_bytes: bytes,
    case_info: dict,
    principals: List[dict],
    agent_photo: Optional[bytes] = None,
) -> tuple:
    """
    Generate a combined POA document for one or more principals.

    principals: list of dicts, each with per-principal fields plus an optional
                "photo" key (bytes) for that principal's ID image.

    Returns: (combined_docx_bytes, total_replacement_count)
    """
    doc_bytes_list = []
    total = 0

    for principal in principals:
        # Merge shared case fields with this principal's fields
        full_info = {**case_info, **{k: v for k, v in principal.items() if k != "photo"}}
        principal_photo = principal.get("photo")

        doc_bytes, count = generate_poa_bytes(
            template_bytes,
            full_info,
            principal_photo=principal_photo,
            agent_photo=agent_photo,
        )
        doc_bytes_list.append(doc_bytes)
        total += count

    combined = _merge_documents(doc_bytes_list)
    return combined, total


# ---------------------------------------------------------------------------
# 2-IP / 3-Agent template support  (Ding & Luo style)
# ---------------------------------------------------------------------------

FIELDS_2IP = [
    # IP 1
    {"key": "ip1_name",     "label": "IP1 Full Name",                  "template_value": "IP1NAME",      "required": True,  "to_upper": True,  "default": ""},
    {"key": "ip1_passport", "label": "IP1 Passport Number",            "template_value": "IP1PASSPORT",  "required": True,  "to_upper": False, "default": ""},
    # IP 2
    {"key": "ip2_name",     "label": "IP2 Full Name",                  "template_value": "IP2NAME",      "required": True,  "to_upper": True,  "default": ""},
    {"key": "ip2_passport", "label": "IP2 Passport Number",            "template_value": "IP2PASSPORT",  "required": True,  "to_upper": False, "default": ""},
    # Case
    {"key": "child_last_name", "label": "Child Last Name",             "template_value": "CHILDNAME",    "required": True,  "to_upper": True,  "default": ""},
    {"key": "surrogate_name",  "label": "Surrogate Full Name",         "template_value": "SURROGATENAME","required": True,  "to_upper": True,  "default": ""},
    {"key": "surrogate_dob",   "label": "Surrogate Date of Birth",     "template_value": "SURROGATEDOB", "required": True,  "to_upper": False, "default": ""},
    {"key": "due_date",        "label": "Due Date",                    "template_value": "DUEDATE",      "required": True,  "to_upper": False, "default": ""},
    {"key": "hospital_name",   "label": "Hospital Name",               "template_value": "HOSPITALNAME", "required": False, "to_upper": False, "default": "Loma Linda University Medical Center"},
    # Agents
    {"key": "agent1_name",  "label": "Agent 1 Full Name",              "template_value": "AGENT1NAME",   "required": True,  "to_upper": True,  "default": ""},
    {"key": "agent1_dob",   "label": "Agent 1 Date of Birth",          "template_value": "AGENT1DOB",    "required": True,  "to_upper": False, "default": ""},
    {"key": "agent1_dl",    "label": "Agent 1 CA Driver License No.",  "template_value": "AGENT1DL",     "required": True,  "to_upper": False, "default": ""},
    {"key": "agent2_name",  "label": "Agent 2 Full Name",              "template_value": "AGENT2NAME",   "required": True,  "to_upper": True,  "default": ""},
    {"key": "agent2_dob",   "label": "Agent 2 Date of Birth",          "template_value": "AGENT2DOB",    "required": True,  "to_upper": False, "default": ""},
    {"key": "agent2_dl",    "label": "Agent 2 CA Driver License No.",  "template_value": "AGENT2DL",     "required": True,  "to_upper": False, "default": ""},
    {"key": "agent3_name",  "label": "Agent 3 Full Name",              "template_value": "AGENT3NAME",   "required": True,  "to_upper": True,  "default": ""},
    {"key": "agent3_dob",   "label": "Agent 3 Date of Birth",          "template_value": "AGENT3DOB",    "required": True,  "to_upper": False, "default": ""},
    {"key": "agent3_dl",    "label": "Agent 3 CA Driver License No.",  "template_value": "AGENT3DL",     "required": True,  "to_upper": False, "default": ""},
    # Firm
    {"key": "attorney_name","label": "Handling Attorney",              "template_value": "ATTORNEYNAME", "required": False, "to_upper": False, "default": "Xuelan Fang"},
    {"key": "agency_name",  "label": "Fertility Agency (for filename)","template_value": "",             "required": False, "to_upper": False, "default": ""},
]

# Longer / more specific strings must come before any substring they contain.
REPLACEMENT_ORDER_2IP = [
    "hospital_name",    # "Loma Linda University Medical Center" — longest
    "surrogate_name",
    "surrogate_dob",
    "due_date",
    "ip1_name",         # IP1NAME before IP1PASSPORT (no actual overlap but keeps intent clear)
    "ip1_passport",
    "ip2_name",
    "ip2_passport",
    "child_last_name",
    "agent1_name",      # AGENT1NAME before AGENT1DOB / AGENT1DL
    "agent1_dob",
    "agent1_dl",
    "agent2_name",
    "agent2_dob",
    "agent2_dl",
    "agent3_name",
    "agent3_dob",
    "agent3_dl",
    "attorney_name",
]

# Image-slot indices (into _find_image_paragraphs result) for the 2IP template.
# Layout: agent1=0, agent2=1, agent3=2, ip1=3+4 (same photo), ip2=5+6 (same photo).
_2IP_PHOTO_SLOTS = {
    "agent1": [0],
    "agent2": [1],
    "agent3": [2],
    "ip1":    [3, 4],
    "ip2":    [5, 6],
}


def generate_2ip_poa_bytes(
    template_bytes: bytes,
    info: dict,
    photos: Optional[dict] = None,
) -> tuple:
    """
    Generate a single 2-IP / 3-agent POA document.

    info: dict with keys matching FIELDS_2IP.
    photos: optional dict with keys "agent1", "agent2", "agent3", "ip1", "ip2" → bytes.
    Returns (docx_bytes, replacement_count).
    """
    doc = Document(io.BytesIO(template_bytes))
    total = 0

    for key in REPLACEMENT_ORDER_2IP:
        field = next((f for f in FIELDS_2IP if f["key"] == key), None)
        if not field:
            continue
        value = info.get(key, "").strip()
        if field["to_upper"]:
            value = value.upper()
        if not value:
            value = field["default"]
        if field["template_value"] and value:
            total += _replace_in_doc(doc, field["template_value"], value)

    if photos:
        img_paras = _find_image_paragraphs(doc)
        for person, slots in _2IP_PHOTO_SLOTS.items():
            photo_bytes = photos.get(person)
            if photo_bytes:
                for slot in slots:
                    if slot < len(img_paras):
                        _replace_image(doc, img_paras[slot], photo_bytes)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), total


def validate_2ip_case(info: dict) -> list:
    """Return list of error strings for a 2IP case. Empty = valid."""
    errors = []
    for field in FIELDS_2IP:
        if field["required"] and not info.get(field["key"], "").strip():
            errors.append(f"'{field['label']}' is required.")
    return errors


def make_filename_2ip(info: dict) -> str:
    ip1_last = (info.get("ip1_name", "IP1") or "IP1").split()[-1]
    ip2_last = (info.get("ip2_name", "IP2") or "IP2").split()[-1]
    sur_last = (info.get("surrogate_name", "Surrogate") or "Surrogate").split()[-1]
    agency = info.get("agency_name", "").strip()
    parts = [f"POA - {ip1_last} & {ip2_last} & {sur_last}"]
    if agency:
        parts.append(agency)
    return " - ".join(parts) + ".docx"


def get_csv_columns() -> list:
    return [f["key"] for f in FIELDS]


def get_example_row() -> dict:
    return {f["key"]: f["template_value"] for f in FIELDS}


def load_cases_from_csv(csv_bytes: bytes) -> tuple:
    text = csv_bytes.decode("utf-8-sig")
    reader = csv.DictReader(io.StringIO(text))
    expected = set(get_csv_columns())
    cases, errors = [], []
    for i, row in enumerate(reader, start=2):
        clean = {k.strip(): v.strip() for k, v in row.items()}
        missing = expected - set(clean.keys())
        if missing:
            errors.append((i, f"Missing columns: {', '.join(sorted(missing))}"))
            continue
        errs = validate_case(clean)
        if errs:
            errors.append((i, "; ".join(errs)))
        else:
            cases.append(clean)
    return cases, errors


# ---------------------------------------------------------------------------
# V2: Unified generator for all 6 combos (1–2 IPs × 1–3 agents)
# ---------------------------------------------------------------------------

COMMON_FIELDS_V2 = [
    {"key": "child_last_name","label": "Child Last Name",          "template_value": "CHILDNAME",    "required": True,  "to_upper": True,  "default": ""},
    {"key": "surrogate_name", "label": "Surrogate Full Name",       "template_value": "SURROGATENAME","required": True,  "to_upper": True,  "default": ""},
    {"key": "surrogate_dob",  "label": "Surrogate Date of Birth",   "template_value": "SURROGATEDOB", "required": True,  "to_upper": False, "default": ""},
    {"key": "due_date",       "label": "Due Date",                  "template_value": "DUEDATE",      "required": True,  "to_upper": False, "default": ""},
    {"key": "hospital_name",  "label": "Hospital Name",             "template_value": "HOSPITALNAME", "required": False, "to_upper": False, "default": "Loma Linda University Medical Center"},
    {"key": "hospital_city",  "label": "Hospital City",             "template_value": "HOSPITALCITY", "required": False, "to_upper": False, "default": "Loma Linda"},
    {"key": "hospital_state", "label": "Hospital State",            "template_value": "HOSPITALSTATE","required": False, "to_upper": False, "default": "California"},
    {"key": "attorney_name",  "label": "Handling Attorney",         "template_value": "ATTORNEYNAME", "required": False, "to_upper": False, "default": "Xuelan Fang"},
    {"key": "agency_name",    "label": "Agency (filename only)",    "template_value": "",             "required": False, "to_upper": False, "default": ""},
]

_IP_FTPL = [
    {"key": "ip{n}_name",    "label": "IP {n} Full Name",        "tval": "IP{N}NAME",    "required": True,  "upper": True,  "default": ""},
    {"key": "ip{n}_role",    "label": "IP {n} Role",             "tval": "IP{N}ROLE",    "required": False, "upper": False, "default": "Intended Father"},
    {"key": "ip{n}_country", "label": "IP {n} Passport Country", "tval": "IP{N}COUNTRY", "required": True,  "upper": False, "default": "People’s Republic of China"},
    {"key": "ip{n}_passport","label": "IP {n} Passport Number",  "tval": "IP{N}PASSPORT","required": True,  "upper": False, "default": ""},
]

_AGENT_FTPL = [
    {"key": "agent{n}_name","label": "Agent {n} Full Name",      "tval": "AGENT{N}NAME","required": True,  "upper": True,  "default": ""},
    {"key": "agent{n}_dob", "label": "Agent {n} Date of Birth",  "tval": "AGENT{N}DOB", "required": True,  "upper": False, "default": ""},
    {"key": "agent{n}_id",  "label": "Agent {n} ID (full phrase)","tval": "AGENT{N}ID",  "required": True,  "upper": False, "default": ""},
]


def get_fields_v2(num_ips: int, num_agents: int) -> list:
    """Return complete field list for the given combo."""
    fields = []
    for n in range(1, num_ips + 1):
        for t in _IP_FTPL:
            fields.append({
                "key":            t["key"].format(n=n),
                "label":          t["label"].format(n=n),
                "template_value": t["tval"].replace("{N}", str(n)),
                "required":       t["required"],
                "to_upper":       t["upper"],
                "default":        t["default"],
            })
    fields.extend(COMMON_FIELDS_V2)
    for n in range(1, num_agents + 1):
        for t in _AGENT_FTPL:
            fields.append({
                "key":            t["key"].format(n=n),
                "label":          t["label"].format(n=n),
                "template_value": t["tval"].replace("{N}", str(n)),
                "required":       t["required"],
                "to_upper":       t["upper"],
                "default":        t["default"],
            })
    return fields


# Photo slot positions (index into _find_image_paragraphs result) for known combos.
_V2_PHOTO_SLOTS = {
    (1, 1): {"ip1": [0], "agent1": [1]},
    (2, 3): {"agent1": [0], "agent2": [1], "agent3": [2], "ip1": [3, 4], "ip2": [5, 6]},
}


def generate_poa_v2_bytes(
    template_bytes: bytes,
    num_ips: int,
    num_agents: int,
    info: dict,
    photos: Optional[dict] = None,
) -> tuple:
    """
    Unified POA generator for all 6 combos.

    info keys:  ip{n}_name, ip{n}_role, ip{n}_country, ip{n}_passport  (n = 1..num_ips)
                child_last_name, surrogate_name, surrogate_dob, due_date,
                hospital_name, hospital_city, hospital_state, attorney_name, agency_name
                agent{n}_name, agent{n}_dob, agent{n}_id              (n = 1..num_agents)

    agent{n}_id = full ID phrase, e.g. "California's Driver License No. Y1234567"

    photos: dict mapping "ip1", "ip2", "agent1", "agent2", "agent3" → bytes (optional).
    Returns (docx_bytes, replacement_count).
    """
    doc = Document(io.BytesIO(template_bytes))
    total = 0

    for field in get_fields_v2(num_ips, num_agents):
        tval = field["template_value"]
        if not tval:
            continue
        value = info.get(field["key"], "").strip()
        if field["to_upper"]:
            value = value.upper()
        if not value:
            value = field["default"]
        if value:
            total += _replace_in_doc(doc, tval, value)

    if photos:
        slots = _V2_PHOTO_SLOTS.get((num_ips, num_agents))
        if slots:
            img_paras = _find_image_paragraphs(doc)
            for person, slot_list in slots.items():
                pbytes = photos.get(person)
                if pbytes:
                    for s in slot_list:
                        if s < len(img_paras):
                            _replace_image(doc, img_paras[s], pbytes)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), total


def validate_v2(num_ips: int, num_agents: int, info: dict) -> list:
    """Return list of error strings. Empty = valid."""
    return [
        f"'{f['label']}' is required."
        for f in get_fields_v2(num_ips, num_agents)
        if f["required"] and not info.get(f["key"], "").strip()
    ]


def make_filename_v2(num_ips: int, info: dict) -> str:
    if num_ips == 1:
        ips = (info.get("ip1_name", "IP1") or "IP1").split()[-1]
    else:
        a = (info.get("ip1_name", "IP1") or "IP1").split()[-1]
        b = (info.get("ip2_name", "IP2") or "IP2").split()[-1]
        ips = f"{a} & {b}"
    sur = (info.get("surrogate_name", "Surrogate") or "Surrogate").split()[-1]
    agency = info.get("agency_name", "").strip()
    parts = [f"POA - {ips} & {sur}"]
    if agency:
        parts.append(agency)
    return " - ".join(parts) + ".docx"
