"""
prep_all_templates.py
Run once to create all 6 poa_template_*.docx files from past case templates.

    python prep_all_templates.py

Outputs (all gitignored / stay local):
    poa_template_1ip_1agent.docx
    poa_template_1ip_2agent.docx
    poa_template_1ip_3agent.docx
    poa_template_2ip_1agent.docx
    poa_template_2ip_2agent.docx
    poa_template_2ip_3agent.docx
"""
import copy
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
OD   = r"C:\Users\zhous\OneDrive - Tsong Law Group\Ralph Tsong's files - Active Cases\Surrogacy Cases"

SOURCES = {
    (1, 1): r"C:\Users\zhous\OneDrive - Tsong Law Group\Ralph Tsong's files - Marketing\POA Pilot\Power of Attorney - Shi & Aispuro - C&T Fertility Consultant.docx",
    (1, 2): os.path.join(OD, r"25-S-178 Zhou & Brummett - Road to Baby\POA\Power of Attorney - Zhou & Brummett - Road to Baby.docx"),
    (1, 3): os.path.join(OD, r"24-S-252 Po-Wei Chen & Maria Blanca Gonzalez - LAS\POA\Power of Attorney - Chen.docx"),
    (2, 1): os.path.join(OD, r"24-s-292 Kwong and Sheng & Pasillas - Pacific Miracles\POA\Power of Attorney - Kwong and Sheng & Pasillas - Pacific Miracles.docx"),
    (2, 2): os.path.join(OD, r"25-S-94 Xinmiao Jia and Zhen Wang & Jessica Lamas - LAS\POA\Power of Attorney - Xinmiao Jia and Zhen Wang & Jessica Lamas - LAS.docx"),
    (2, 3): r"C:\Users\zhous\OneDrive - Tsong Law Group\Ralph Tsong's files - Marketing\POA Pilot\Power of Attorney - Ding and Luo & Sanchez-Diaz - Lily Baby Surrogacy.docx",
}

OUTPUTS = {k: os.path.join(HERE, f"poa_template_{k[0]}ip_{k[1]}agent.docx") for k in SOURCES}

R = "’"   # right single quote '
LD = "“"  # left double quote "
RD = "”"  # right double quote "


# ---------------------------------------------------------------------------
# Paragraph text builders
# ---------------------------------------------------------------------------

def _agents(n, conj="and/or"):
    names = [f"AGENT{i}NAME" for i in range(1, n + 1)]
    if n == 1:   return names[0]
    if n == 2:   return f"{names[0]} {conj} {names[1]}"
    return f"{names[0]}, {names[1]}, {conj} {names[2]}"

def _consent_agents(n):
    """For para 15 prefix: 'AGENT1NAME, AGENT2NAME' (no trailing and/or)"""
    names = [f"AGENT{i}NAME" for i in range(1, n + 1)]
    return ", ".join(names)

def p7(nip):
    if nip == 1: return "Attorney for IP1NAME"
    return "Attorney for IP1NAME and IP2NAME"

def p10(nip):
    if nip == 1:
        return (
            f"I, IP1NAME, am the IP1ROLE of the Infant CHILDNAME, born to me via gestational "
            f"surrogate SURROGATENAME. My IP1COUNTRY Passport No. is IP1PASSPORT.  "
            f"A photocopy of my passport is attached to this declaration."
        )
    return (
        f"We, IP1NAME AND IP2NAME, are the Intended Parents of the Infant CHILDNAME, born to us via "
        f"gestational surrogate SURROGATENAME. IP1NAME{R}s IP1COUNTRY Passport No. is IP1PASSPORT, "
        f"and IP2NAME{R}s IP2COUNTRY Passport No. is IP2PASSPORT.  "
        f"Photocopies of our passports are attached to this declaration."
    )

def p11(nip):
    poss = "My" if nip == 1 else "Our"
    return (
        f"{poss} child, Infant CHILDNAME is expected to be born anytime from now to the original "
        f"anticipated due date, on or before DUEDATE, at HOSPITALNAME, in HOSPITALCITY, HOSPITALSTATE "
        f"via gestational surrogate SURROGATENAME, whose date of birth is SURROGATEDOB."
    )

def p12(nip, na):
    iw = "I" if nip == 1 else "We"
    def entry(n):
        return (
            f"AGENT{n}NAME, date of birth AGENT{n}DOB and AGENT{n}ID "
            f"(a copy of AGENT{n}NAME{R}s identification is attached to this declaration)"
        )
    if na == 1:
        return f"{iw} hereby consent to giving special powers of attorney to {entry(1)}."
    if na == 2:
        return (
            f"{iw} hereby consent to giving special powers of attorney to "
            f"{entry(1)}, and {entry(2)}."
        )
    return (
        f"{iw} hereby consent to giving special powers of attorney to "
        f"{entry(1)}, {entry(2)}, and {entry(3)}."
    )

def p13(nip, na):
    iw    = "I" if nip == 1 else "We"
    mous  = "me" if nip == 1 else "us"
    myour = "my" if nip == 1 else "our"
    if na == 1:
        astr = "AGENT1NAME is to be provided a wristband, allowing access"
    elif na == 2:
        astr = f"AGENT1NAME and/or AGENT2NAME are to be provided a wristband, allowing their access"
    else:
        astr = f"AGENT1NAME, AGENT2NAME, and/or AGENT3NAME are to be provided wristbands, allowing them access"
    return (
        f"{iw} understand and agree that {astr} to the nursery "
        f"(including the Neo-Natal Intensive Care Unit {LD}NICU{RD}) by any physician and/or medical "
        f"facility and/or medical personnel providing post-natal or well care to the child born to "
        f"{mous} by {myour} gestational surrogate."
    )

def p14(nip, na):
    iw   = "I" if nip == 1 else "We"
    mous = "my" if nip == 1 else "our"
    fact = "I am" if nip == 1 else "we are"
    if na == 1:   discharge = "AGENT1NAME"
    elif na == 2: discharge = "AGENT1NAME or AGENT2NAME"
    else:         discharge = "AGENT1NAME, AGENT2NAME, or AGENT3NAME"
    return (
        f"{iw} hereby consent that the child born to {mous.replace('my','me').replace('our','us')} by "
        f"SURROGATENAME shall be released upon discharge to {discharge} without {mous} presence. "
        f"This is due to the fact {fact} not able to arrive the United States before Child{R}s delivery."
    )

def p15(nip, na):
    iw    = "I" if nip == 1 else "We"
    myour = "my" if nip == 1 else "our"
    agents = _consent_agents(na)
    prefix = f"{agents} and/or"
    return (
        f"{iw} hereby consent for {prefix} any agent designated by {myour} attorney ATTORNEYNAME "
        f"to complete the Birth Certificate Application and sign the birth certificate for {myour} child "
        f"so that no paperwork is delayed in having the Certified Birth Certificate available. "
        f"{iw} give permission for {prefix} any agent designated by {myour} attorney ATTORNEYNAME "
        f"to provide the full name that {iw.lower()} wish to give {myour} child."
    )

def p16(nip, na):
    iw    = "I" if nip == 1 else "We"
    myour = "my" if nip == 1 else "our"
    meus  = "me" if nip == 1 else "us"
    if na == 1:
        astr = "AGENT1NAME"
        is_are = "is"
        rep_agents = "AGENT1NAME"
        auth_agents = "AGENT1NAME"
    elif na == 2:
        astr = "AGENT1NAME and/or AGENT2NAME"
        is_are = "are"
        rep_agents = astr
        auth_agents = astr
    else:
        astr = "AGENT1NAME, AGENT2NAME, and/or AGENT3NAME"
        is_are = "are"
        rep_agents = astr
        auth_agents = astr
    return (
        f"{iw} hereby consent for {astr} to act as true and lawful representative, agent, and "
        f"Attorney-In-Fact for {meus} and in their name to sign documents on {myour} behalf and to act "
        f"on {myour} behalf as needed for the hospital and further pediatric well care. "
        f"{iw} consent and authorize for {rep_agents} to be able to speak with the hospital staff "
        f"regarding the health status and to receive medical information. "
        f"{auth_agents} {is_are} authorized to make decisions for the benefit of and best interest of "
        f"{myour} child after birth. {iw} understand that additional documents may be required in order "
        f"to effectuate such powers due to the requirements of third parties."
    )

def p17(nip, na):
    iw = "I" if nip == 1 else "We"
    if na == 1:   astr = "AGENT1NAME"
    elif na == 2: astr = "AGENT1NAME and/or AGENT2NAME"
    else:         astr = "AGENT1NAME, AGENT2NAME, or AGENT3NAME"
    return (
        f"{iw} hereby authorize {astr} to obtain a passport for Infant CHILDNAME and to travel with "
        f"the Infant CHILDNAME including to any state or outside of the United States."
    )

def p18(nip, na):
    iw    = "I" if nip == 1 else "We"
    myour = "my" if nip == 1 else "our"
    if na == 1:   astr = "AGENT1NAME"
    elif na == 2: astr = "AGENT1NAME and/or AGENT2NAME"
    else:         astr = "AGENT1NAME, AGENT2NAME, or AGENT3NAME"
    return (
        f"{iw} hereby authorize {astr} to obtain Infant CHILDNAME{R}s birth certificate, travel "
        f"documents, and to request an apostille from any applicable agency on {myour} behalf."
    )

def p21(nip):
    if nip == 1:
        return "IN WITNESS WHEREOF, I, IP1NAME, have executed this Power of Attorney on ______________."
    return "IN WITNESS WHEREOF, We, IP1NAME AND IP2NAME, have executed this Power of Attorney on ______________."

def p23(nip):
    if nip == 1: return "IP1NAME"
    return "IP1NAME AND IP2NAME"

def cell_para4(nip):
    if nip == 1: return "to be born to parent IP1NAME"
    return "to be born to parentS IP1NAME AND IP2NAME"


# ---------------------------------------------------------------------------
# Low-level: replace all run content in a paragraph
# ---------------------------------------------------------------------------

def set_para_text(para, new_text: str) -> None:
    p = para._p
    first_r = p.find(qn("w:r"))
    first_rpr = None
    if first_r is not None:
        rpr = first_r.find(qn("w:rPr"))
        if rpr is not None:
            first_rpr = copy.deepcopy(rpr)
    for r in p.findall(qn("w:r")):
        p.remove(r)
    for i, seg in enumerate(new_text.split("\t")):
        if i > 0:
            tab_r = OxmlElement("w:r")
            if first_rpr is not None:
                tab_r.append(copy.deepcopy(first_rpr))
            tab_r.append(OxmlElement("w:tab"))
            p.append(tab_r)
        if seg:
            text_r = OxmlElement("w:r")
            if first_rpr is not None:
                text_r.append(copy.deepcopy(first_rpr))
            t = OxmlElement("w:t")
            t.text = seg
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            text_r.append(t)
            p.append(text_r)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def prep(nip, na):
    src = SOURCES[(nip, na)]
    out = OUTPUTS[(nip, na)]

    if not os.path.exists(src):
        print(f"  SKIP ({nip}IP/{na}A) — source not found: {os.path.basename(src)}")
        return

    doc = Document(src)
    paras = doc.paragraphs

    replacements = {
        7:  p7(nip),
        10: p10(nip),
        11: p11(nip),
        12: p12(nip, na),
        13: p13(nip, na),
        14: p14(nip, na),
        15: p15(nip, na),
        16: p16(nip, na),
        17: p17(nip, na),
        18: p18(nip, na),
        21: p21(nip),
        23: p23(nip),
    }

    for idx, text in sorted(replacements.items()):
        if idx < len(paras):
            set_para_text(paras[idx], text)

    # Fix table caption cell
    try:
        cell = doc.tables[0].rows[0].cells[0]
        if len(cell.paragraphs) > 2:
            set_para_text(cell.paragraphs[2], "       infant CHILDNAME")
        if len(cell.paragraphs) > 4:
            set_para_text(cell.paragraphs[4], cell_para4(nip))
    except (IndexError, AttributeError):
        pass

    doc.save(out)
    print(f"  OK  ({nip}IP/{na}A) -> {os.path.basename(out)}")


def main():
    print("Prepping all templates...\n")
    for (nip, na) in [(1,1),(1,2),(1,3),(2,1),(2,2),(2,3)]:
        prep(nip, na)
    print("\nDone. Upload each poa_template_*.docx via the app sidebar.")


if __name__ == "__main__":
    main()
