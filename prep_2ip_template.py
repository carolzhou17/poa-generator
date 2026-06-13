"""
prep_2ip_template.py
One-time script: inserts placeholder text into the Ding & Luo template
so poa_generator can fill it in automatically.

Run from the "PBO auto" folder:
    python prep_2ip_template.py

Output: poa_template_2ip.docx  (stays local — listed in .gitignore)
"""
import copy
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SOURCE = (
    r"C:\Users\zhous\OneDrive - Tsong Law Group"
    r"\Ralph Tsong's files - Marketing\POA Pilot"
    r"\Power of Attorney - Ding and Luo & Sanchez-Diaz - Lily Baby Surrogacy.docx"
)
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "poa_template_2ip.docx")

RSQUO = "’"   # '
LDQUO = "“"   # "
RDQUO = "”"   # "

# Maps paragraph index → replacement text (placeholders inserted)
PARA_TEXT = {
    7: "Attorney for IP1NAME and IP2NAME",
    10: (
        f"We, IP1NAME AND IP2NAME, are the Intended Parents of the Infant CHILDNAME, "
        f"born to us via gestational surrogate SURROGATENAME. "
        f"IP1NAME{RSQUO}s People{RSQUO}s Republic of China Passport No. is IP1PASSPORT, "
        f"and IP2NAME{RSQUO}s People{RSQUO}s Republic of China Passport No. is IP2PASSPORT.  "
        f"Photocopies of our passports are attached to this declaration."
    ),
    11: (
        f"Our child, Infant CHILDNAME is expected to be born anytime from now to the "
        f"original anticipated due date, on or before DUEDATE, at HOSPITALNAME via "
        f"gestational surrogate SURROGATENAME, whose date of birth is SURROGATEDOB."
    ),
    12: (
        f"We hereby consent to giving special powers of attorney to AGENT1NAME, date of birth "
        f"AGENT1DOB and California{RSQUO}s Driver License No. AGENT1DL "
        f"(a copy of AGENT1NAME{RSQUO}s identification is attached to this declaration), "
        f"AGENT2NAME, date of birth AGENT2DOB and California{RSQUO}s Driver License No. AGENT2DL "
        f"(a copy of AGENT2NAME{RSQUO}s identification is attached to this declaration), "
        f"and AGENT3NAME, date of birth AGENT3DOB and California{RSQUO}s Driver License No. AGENT3DL "
        f"(a copy of AGENT3NAME{RSQUO}s identification is attached to this declaration)."
    ),
    13: (
        f"We understand and agree that AGENT1NAME, AGENT2NAME, and/or AGENT3NAME are to be "
        f"provided a wristband, allowing their access to the nursery (including the Neo-Natal "
        f"Intensive Care Unit {LDQUO}NICU{RDQUO}) by any physician and/or medical facility and/or "
        f"medical personnel providing post-natal or well care to the child born to us by our "
        f"gestational surrogate."
    ),
    14: (
        f"We hereby consent that the child born to us by SURROGATENAME shall be released upon "
        f"discharge to AGENT1NAME, AGENT2NAME, and/or AGENT3NAME without our presence. This is "
        f"due to the fact we are not able to arrive the United States before Child{RSQUO}s delivery."
    ),
    15: (
        f"We hereby consent for AGENT1NAME, AGENT2NAME, and/or AGENT3NAME and/or any agent "
        f"designated by our attorney ATTORNEYNAME to complete the Birth Certificate Application "
        f"and sign the birth certificate for our child so that no paperwork is delayed in having "
        f"the Certified Birth Certificate available. We give permission for AGENT1NAME, AGENT2NAME, "
        f"and/or AGENT3NAME and/or any agent designated by our attorney ATTORNEYNAME to provide "
        f"the full name that we wish to give our child."
    ),
    16: (
        f"We hereby consent for AGENT1NAME, AGENT2NAME, and/or AGENT3NAME to act as true and "
        f"lawful representative, agent, and Attorney-In-Fact for us and in their name to sign "
        f"documents on our behalf and to act on our behalf as needed for the hospital and further "
        f"pediatric well care. We consent and authorize for AGENT1NAME, AGENT2NAME, and/or "
        f"AGENT3NAME to be able to speak with the hospital staff regarding the health status and "
        f"to receive medical information. AGENT1NAME, AGENT2NAME, and/or AGENT3NAME is authorized "
        f"to make decisions for the benefit of and best interest of our child after birth. We "
        f"understand that additional documents may be required in order to effectuate such powers "
        f"due to the requirements of third parties."
    ),
    17: (
        f"We hereby authorize for AGENT1NAME, AGENT2NAME, and/or AGENT3NAME to obtain a passport "
        f"for Infant CHILDNAME and to travel with the Infant CHILDNAME including to any state or "
        f"outside of the United States."
    ),
    18: (
        f"We hereby authorize AGENT1NAME, AGENT2NAME, and/or AGENT3NAME to obtain Infant "
        f"CHILDNAME{RSQUO}s birth certificate, travel documents, and to request an apostille from "
        f"any applicable agency on our behalf."
    ),
    21: (
        "\tIN WITNESS WHEREOF, We, IP1NAME AND IP2NAME, have executed this "
        "Power of Attorney on ______________."
    ),
}


def set_para_text(para, new_text: str) -> None:
    """Replace all run content in a paragraph with new_text, preserving paragraph formatting."""
    p = para._p

    # Save run character properties from the first run (bold, font, size, etc.)
    first_r = p.find(qn("w:r"))
    first_rpr = None
    if first_r is not None:
        rpr = first_r.find(qn("w:rPr"))
        if rpr is not None:
            first_rpr = copy.deepcopy(rpr)

    # Remove all existing runs
    for r in p.findall(qn("w:r")):
        p.remove(r)

    # Rebuild runs, splitting on tab characters
    segments = new_text.split("\t")
    for i, segment in enumerate(segments):
        if i > 0:
            tab_r = OxmlElement("w:r")
            if first_rpr is not None:
                tab_r.append(copy.deepcopy(first_rpr))
            tab_r.append(OxmlElement("w:tab"))
            p.append(tab_r)
        if segment:
            text_r = OxmlElement("w:r")
            if first_rpr is not None:
                text_r.append(copy.deepcopy(first_rpr))
            t = OxmlElement("w:t")
            t.text = segment
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            text_r.append(t)
            p.append(text_r)


def main():
    if not os.path.exists(SOURCE):
        print(f"ERROR: Source template not found:\n  {SOURCE}")
        sys.exit(1)

    doc = Document(SOURCE)

    # Replace body paragraphs
    for idx in sorted(PARA_TEXT.keys()):
        set_para_text(doc.paragraphs[idx], PARA_TEXT[idx])
        preview = PARA_TEXT[idx].replace("\t", "\\t")[:70]
        print(f"  [para {idx:2d}] {preview!r}")

    # Fix the case-caption table (first table, first cell)
    cell = doc.tables[0].rows[0].cells[0]
    set_para_text(cell.paragraphs[2], "       infant CHILDNAME")
    print("  [table cell-para 2] '       infant CHILDNAME'")
    set_para_text(cell.paragraphs[4], "to be born to parentS IP1NAME AND IP2NAME")
    print("  [table cell-para 4] 'to be born to parentS IP1NAME AND IP2NAME'")

    doc.save(OUTPUT)
    print(f"\n  Saved: {OUTPUT}")
    print("  Upload this file as the '2 IPs + 3 Agents' template in the sidebar.")


if __name__ == "__main__":
    main()
